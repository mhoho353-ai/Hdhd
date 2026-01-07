import streamlit as st
import whisper
from docx import Document
from docx.shared import Pt
import tempfile
import os

st.set_page_config(page_title="تفريغ صوتي احترافي", layout="centered")
st.title("🎧 تفريغ المحاضرات الصوتية (عربي)")

@st.cache_resource
def load_model():
    return whisper.load_model("medium")

model = load_model()

def clean_text(text):
    text = text.replace("  ", " ")
    text = text.replace("،", "، ")
    text = text.replace(".", ".\n")
    text = text.replace("؟", "؟\n")
    return text.strip()

audio = st.file_uploader("ارفع ملف صوتي", type=["mp3","wav","m4a"])

if audio:
    st.success("تم رفع الملف")

    if st.button("ابدأ التفريغ"):
        with st.spinner("جارى التفريغ..."):
            with tempfile.NamedTemporaryFile(delete=False) as tmp:
                tmp.write(audio.read())
                audio_path = tmp.name

            result = model.transcribe(audio_path, language="ar", fp16=False)
            segments = result["segments"]

            final_text = ""
            for seg in segments:
                final_text += clean_text(seg["text"]) + "\n\n"

            doc = Document()
            doc.add_heading("تفريغ محاضرة صوتية", level=1)
            p = doc.add_paragraph(final_text)
            for run in p.runs:
                run.font.size = Pt(14)

            file_name = "transcription.docx"
            doc.save(file_name)

            st.success("تم الانتهاء بنجاح ✅")
            st.download_button(
                "تحميل ملف Word",
                open(file_name, "rb"),
                file_name=file_name
            )

            os.remove(audio_path)
